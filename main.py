from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time

# --- 1. Fungsi Tambahan: Bikin Daftar Isi Otomatis --- 
def tambah_daftar_isi_otomatis(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    # Bikin elemen XML untuk field 'TOC'
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # Perintah ke Word: "Bikin TOC dari Heading 1 sampai 3, kasih link, kasih nomor hal"
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:tldCharType'), 'end')
    
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)

# --- 2. Simulasi AI ---
def panggil_ai_paraphrase(teks_asli):
    if len(teks_asli) < 20: return teks_asli
    return f"[AI-REWRITE] {teks_asli} (Lolos Turnitin)"

# --- 3. Deteksi Judul ---
def cek_apakah_judul(paragraph):
    teks = paragraph.text.strip().upper()
    style_name = paragraph.style.name
    
    # Cek Style
    if style_name.startswith('Heading') or style_name == 'Title':
        return True
        
    # Cek Kata Kunci
    kata_kunci_judul = ["BAB ", "DAFTAR ISI", "DAFTAR PUSTAKA", "KATA PENGANTAR", "ABSTRAK", "PENDAHULUAN", "LATAR BELAKANG"]
    if any(k in teks for k in kata_kunci_judul) and len(teks.split()) < 10:
        return True
        
    return False

# --- 4. Proses Paraphrase ---
def proses_paraphrase(doc):
    print("   -> Sedang membaca & menulis ulang paragraf...")
    for para in doc.paragraphs:
        teks_lama = para.text.strip()
        if not teks_lama: continue
            
        if cek_apakah_judul(para):
            print(f"      [SKIP JUDUL] {teks_lama[:30]}...")
            continue 
            
        teks_baru = panggil_ai_paraphrase(teks_lama)
        para.text = teks_baru

# --- 5. Formatting Margin ---
def atur_margin(doc):
    print("   -> Mengatur margin 4-4-3-3...")
    for section in doc.sections:
        section.top_margin = Cm(4)
        section.left_margin = Cm(4)
        section.bottom_margin = Cm(3)
        section.right_margin = Cm(3)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

# --- 6. BUAT FILE DUMMY (YANG DIPERBAIKI) ---
def buat_file_dummy(nama_file):
    doc = Document()
    
    doc.add_heading('DAFTAR ISI', level=1)

    tambah_daftar_isi_otomatis(doc)

    doc.add_page_break()
    
    # BAB 1
    doc.add_heading('BAB 1 PENDAHULUAN', level=1) 
    doc.add_paragraph('Ini adalah paragraf pembuka.')
    
    # Sub-bab
    doc.add_heading('1.1 Latar Belakang', level=2)
    doc.add_paragraph('Penelitian ini sangat penting.')

    # Sub-bab lagi
    doc.add_heading('1.2 Rumusan Masalah', level=2)
    doc.add_paragraph('Masalahnya adalah skripsi susah.')

    doc.add_page_break()

    # DAFTAR PUSTAKA
    doc.add_heading('DAFTAR PUSTAKA', level=1)
    doc.add_paragraph('Google, 2024.')
    
    doc.save(nama_file)

# --- EKSEKUSI ---
if __name__ == "__main__":
    file_masuk = 'daftar_isi.docx'
    file_keluar = 'hasil_siap_sidang.docx'
    
    try:
        buat_file_dummy(file_masuk)
        
        print("=== MULAI PROSES ===")
        document = Document(file_masuk)
        
        proses_paraphrase(document)
        atur_margin(document)
        
        document.save(file_keluar)
        print(f"=== SELESAI ===\nCek file: {file_keluar}")
        print("NOTE: Saat buka file di Word, jika Daftar Isi belum muncul angka halamannya,")
        print("klik kanan di area Daftar Isi -> Update Field -> Update entire table.")
        
    except PermissionError:
        print("ERROR: Tutup dulu file Word-nya woy! :D")