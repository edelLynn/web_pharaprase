import streamlit as st
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from groq import Groq
from docxcompose.composer import Composer
import os

if "GROQ_API_KEY" in st.secrets:
    API_KEY = st.secrets["GROQ_API_KEY"]
else:
    API_KEY = ""

# Setup Client Groq
client = None
if API_KEY.startswith("gsk_"):
    client = Groq(api_key=API_KEY)

client = None
if API_KEY.startswith("gsk_") and "XXX" not in API_KEY:
    client = Groq(api_key=API_KEY)

# Fungsi 1: Isi Template Cover
def isi_template_cover(path_template, data_input):
    doc = Document(path_template)
    replacements = {
        "[JUDUL_MAKALAH]": data_input['judul'].upper(),
        "[MATA_KULIAH]": data_input['matkul'],
        "[NAMA_MAHASISWA]": data_input['anggota'],
        "[NIM_MAHASISWA]": "", 
        "[NAMA_DOSEN_1]": data_input['dosen1'],
        "[NAMA_DOSEN_2]": data_input['dosen2'] if data_input['dosen2'] else " ",
        "[NAMA_FAKULTAS]": data_input['fakultas'].upper(),
        "[TAHUN]": data_input['tahun']
    }

    for para in doc.paragraphs:
        for kode, isi_baru in replacements.items():
            if kode in para.text:
                para.text = para.text.replace(kode, isi_baru)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.color.rgb = RGBColor(0,0,0)
                    run.bold = True
                    run.font.size = Pt(14 if kode == "[JUDUL_MAKALAH]" else 12)
    return doc

# Fungsi 2: Paraphrase AI
def panggil_ai_paraphrase(teks_asli):
    if not client: return f"[GAGAL API KEY]"
    if len(teks_asli) < 30: return teks_asli

    try:
        chat_completion = client.chat.completions.create(
            messages=[
                {"role": "system", "content": "You are a direct text rewriting machine. Output ONLY the rewritten text in Indonesian academic style. Do NOT provide options. Do NOT change proper names/dates."},
                {"role": "user", "content": f"Rewrite formally: '{teks_asli}'"}
            ],
            model="llama-3.3-70b-versatile", 
            temperature=0.5,
        )
        hasil = chat_completion.choices[0].message.content.replace('"', '').strip()
        if "\n" in hasil: hasil = hasil.split("\n")[0]
        return hasil
    except:
        return teks_asli

# Fungsi 3: Deteksi Judul
def proses_judul(paragraph):
    teks = paragraph.text.strip().upper()
    
    if len(teks.split()) < 15 and "." in teks[:6] and any(c.isdigit() for c in teks[:3]):
        paragraph.style = 'Heading 2'
        paragraph.paragraph_format.page_break_before = False 
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = 'Times New Roman'
            run.bold = True
            run.font.size = Pt(12)
        return True 

    kunci_awal = ["BAB ", "DAFTAR ISI", "KATA PENGANTAR", "ABSTRAK", "PENDAHULUAN", "KESIMPULAN", "SARAN", "PENUTUP", "HASIL DAN PEMBAHASAN"]
    
    is_heading_1 = False
    for k in kunci_awal:
        if teks.startswith(k):
            is_heading_1 = True
            break
            
    if is_heading_1 and len(teks.split()) < 15:
        paragraph.style = 'Heading 1'
        paragraph.paragraph_format.page_break_before = True 
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = 'Times New Roman'
            run.bold = True
            run.font.size = Pt(14)
        return True

    return False

# Fungsi 4: Bersihkan Baris Kosong di Awal Dokumen
def bersihkan_awal_dokumen(doc):
    while len(doc.paragraphs) > 0:
        p = doc.paragraphs[0]
        if not p.text.strip(): 
            p._element.getparent().remove(p._element)
        else:
            break
    return doc

# Fungsi 5: Proses Isi Dokumen
def proses_konten_user(doc):
    doc = bersihkan_awal_dokumen(doc)

    for section in doc.sections:
        section.top_margin = Cm(4); section.left_margin = Cm(4)
        section.bottom_margin = Cm(3); section.right_margin = Cm(3)
        section.page_width = Cm(21); section.page_height = Cm(29.7)

    list_para = [p for p in doc.paragraphs if p.text.strip()]
    total = len(list_para)
    my_bar = st.progress(0, text="Sedang memproses isi...")

    for i, para in enumerate(list_para):
        teks_lama = para.text.strip()
        persen = int(((i+1)/total)*100)
        if persen > 100: persen = 100
        my_bar.progress(persen, text=f"Processing: {teks_lama[:30]}...")

        if proses_judul(para):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if para.style.name == 'Heading 1' else WD_PARAGRAPH_ALIGNMENT.LEFT
            continue

        if "DAFTAR PUSTAKA" in teks_lama.upper():
            para.clear()
            continue

        para.text = panggil_ai_paraphrase(teks_lama)
        
        para.paragraph_format.page_break_before = False
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        para.style = doc.styles['Normal']
        for run in para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = False
    
    my_bar.empty()
    return doc

# Fungsi 6: Tambah Daftar Pustaka
def tambah_daftar_pustaka(doc, teks_dapus):
    if not teks_dapus.strip(): return doc
    doc.add_page_break()
    p_judul = doc.add_paragraph("DAFTAR PUSTAKA")
    p_judul.style = 'Heading 1'
    p_judul.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in p_judul.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0,0,0)
        run.bold = True
        
    list_item = teks_dapus.split('\n')
    for item in list_item:
        if item.strip():
            p = doc.add_paragraph(item.strip())
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0,0,0)
    return doc

# Fungsi 7: Gabung
def gabung_dokumen(doc_cover, doc_isi):
    doc_cover.add_page_break()
    composer = Composer(doc_cover)
    composer.append(doc_isi)
    return composer.doc


st.set_page_config(page_title="Skripsi Generator", page_icon="🎓")
st.title("🎓 Skripsi Auto-Format & Cover Generator")

with st.sidebar:
    st.header("📝 Data Mahasiswa")
    in_judul = st.text_area("Judul Makalah", "")
    in_matkul = st.text_input("Mata Kuliah", "")
    in_anggota = st.text_area("Nama & NIM Anggota (Enter per baris)", 
                              "Nama Ketua - NIM\nNama Anggota 1 - NIM", height=100)
    in_fakultas = st.text_input("Fakultas", "")
    in_tahun = st.text_input("Tahun", "")
    
    st.write("👨‍🏫 **Data Dosen**")
    in_dosen1 = st.text_input("Dosen Pengampu 1", "")
    in_dosen2 = st.text_input("Dosen Pengampu 2", "")

    st.markdown("---")
    st.header("📚 Daftar Pustaka")
    st.info("Paste daftar pustaka di sini.")
    in_dapus = st.text_area("Input Daftar Pustaka", height=200)

kampus_options = ["Universitas Jember (UNEJ)", "Tanpa Template"]
pilihan_kampus = st.selectbox("🏫 Pilih Template Kampus:", kampus_options)

if not client:
    st.error("⚠️ API KEY ERROR! Cek file app.py.")
else:
    st.write("---")
    st.warning("📄 **PENTING:** Upload isi makalah saja (Kata Pengantar, Bab 1, dst).")
    uploaded_file = st.file_uploader("Upload Isi Makalah (.docx)", type="docx")

    if uploaded_file and st.button("🚀 GENERATE SKRIPSI FULL"):
        temp_input = "temp_isi.docx"
        temp_output = "hasil_full.docx"
        with open(temp_input, "wb") as f: f.write(uploaded_file.getbuffer())
        
        try:
            data_user = {
                'judul': in_judul, 'matkul': in_matkul, 
                'anggota': in_anggota,
                'nim': "",
                'fakultas': in_fakultas, 'tahun': in_tahun,
                'dosen1': in_dosen1, 'dosen2': in_dosen2
            }

            st.write("⚙️ Memproses isi makalah...")
            doc_isi = Document(temp_input)
            doc_isi_processed = proses_konten_user(doc_isi)
            
            if in_dapus.strip():
                st.write("📚 Menambahkan Daftar Pustaka...")
                doc_isi_processed = tambah_daftar_pustaka(doc_isi_processed, in_dapus)

            doc_final = None
            if pilihan_kampus == "Universitas Jember (UNEJ)":
                path_cover = os.path.join("cover", "cover_unej.docx")
                if os.path.exists(path_cover):
                    st.write("📄 Menggabungkan Cover...")
                    doc_cover_filled = isi_template_cover(path_cover, data_user)
                    doc_final = gabung_dokumen(doc_cover_filled, doc_isi_processed)
                else:
                    st.error("❌ File cover_unej.docx hilang!")
                    doc_final = doc_isi_processed
            else:
                doc_final = doc_isi_processed

            doc_final.save(temp_output)
            st.balloons()
            st.success("✅ SELESAI!")
            
            with open(temp_output, "rb") as file:
                st.download_button("📥 Download Makalah Jadi", file, "Makalah_Final.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                
        except Exception as e:
            st.error(f"Error: {e}")
